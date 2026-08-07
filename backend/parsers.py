"""
Parses PDF / DOCX / RTF byte content into a flat list of paragraph-level
text blocks, tagged with an optional page number (PDF only).
"""

import io
import logging
from dataclasses import dataclass
from typing import List, Optional, Union

import pdfplumber
from bs4 import BeautifulSoup
from docx import Document
from striprtf.striprtf import rtf_to_text

logger = logging.getLogger("nku_extractor.parsers")


@dataclass
class TextBlock:
    text: str
    page: Optional[int] = None  # 1-indexed page number for PDFs, None otherwise


def parse_pdf(data: bytes) -> List[TextBlock]:
    blocks: List[TextBlock] = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if not text.strip():
                    continue
                # Split on blank lines to approximate paragraph breaks.
                for chunk in text.split("\n\n"):
                    chunk = chunk.strip()
                    if chunk:
                        blocks.append(TextBlock(text=chunk, page=page_num))
    except Exception as exc:
        logger.error("Failed to parse PDF: %s", exc)
    return blocks


def parse_docx(data: bytes) -> List[TextBlock]:
    blocks: List[TextBlock] = []
    try:
        doc = Document(io.BytesIO(data))
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                blocks.append(TextBlock(text=text, page=None))
        # Tables often carry substantive content in these documents too.
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    blocks.append(TextBlock(text=row_text, page=None))
    except Exception as exc:
        logger.error("Failed to parse DOCX: %s", exc)
    return blocks


def parse_rtf(data: bytes) -> List[TextBlock]:
    blocks: List[TextBlock] = []
    try:
        # RTF is text-based but may carry an explicit encoding; try utf-8,
        # fall back to cp1250 (common for older Slovak/Czech RTF exports).
        try:
            raw = data.decode("utf-8")
        except UnicodeDecodeError:
            raw = data.decode("cp1250", errors="replace")
        plain = rtf_to_text(raw)
        for chunk in plain.split("\n\n"):
            chunk = chunk.strip()
            if chunk:
                blocks.append(TextBlock(text=chunk, page=None))
    except Exception as exc:
        logger.error("Failed to parse RTF: %s", exc)
    return blocks


def parse_document_preview_html(html: Union[str, bytes]) -> List[TextBlock]:
    """
    Parses the HTML returned by nrsr.sk's Dynamic/DocumentPreview.aspx endpoint.

    VERIFIED 2026-07-22 against a live response (DocID=577477, 67 pages):
    - This is NOT a downloadable file and there is no separate Download.aspx
      link behind it -- the full document text is server-rendered directly
      into the HTML (Aspose.Words-style export), one <div class="awdiv awpage">
      per page, containing many absolutely-positioned <span> word/run elements.
    - The ENTIRE document (all pages) comes back in a single HTTP response --
      there is no per-page pagination via separate requests. A 67-page
      document was ~4.3MB of HTML in one GET.
    - It is static HTML, not JS-rendered, so plain BeautifulSoup text
      extraction works; no headless browser is needed.

    Each awpage div is treated as one TextBlock with page = its 1-indexed
    position, matching the page_hint semantics already used for PDFs.
    """
    blocks: List[TextBlock] = []
    try:
        soup = BeautifulSoup(html, "html.parser")
        pages = soup.find_all("div", class_="awpage")

        if not pages:
            # Markup may have changed since verification -- degrade gracefully
            # rather than silently returning nothing.
            logger.warning(
                "No 'awpage' divs found in DocumentPreview HTML -- markup may "
                "have changed since this was last verified (2026-07-22). "
                "Falling back to whole-body text extraction (page numbers "
                "will be unavailable)."
            )
            body_text = soup.get_text(" ", strip=True)
            if body_text:
                blocks.append(TextBlock(text=body_text, page=None))
            return blocks

        for page_num, page in enumerate(pages, start=1):
            text = page.get_text(" ", strip=True)
            if text:
                blocks.append(TextBlock(text=text, page=page_num))
    except Exception as exc:
        logger.error("Failed to parse DocumentPreview HTML: %s", exc)
    return blocks


def parse_document(data: bytes, filename: str) -> List[TextBlock]:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return parse_pdf(data)
    if lower.endswith(".docx"):
        return parse_docx(data)
    if lower.endswith(".doc"):
        # Legacy binary .doc is not supported by python-docx (XML-based).
        # Fail loudly rather than silently returning garbage.
        logger.warning(
            "Legacy .doc (binary) format detected for %s -- not supported. "
            "Convert to .docx/.pdf, or add antiword/LibreOffice conversion.",
            filename,
        )
        return []
    if lower.endswith(".rtf"):
        return parse_rtf(data)
    logger.warning("Unsupported file extension for %s", filename)
    return []
